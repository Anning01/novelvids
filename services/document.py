"""书稿正文提取与章节结构校验（上传、OSS 终局、项目创建共用）。

将纯文本/ Markdown / DOCX / PDF 的正文提取与章节质量校验收敛到一处，
供 `api/file.py`（上传与 OSS 终局）与 `controllers/novel.py`（经 OSS key 创建项目）复用，
避免重复解析逻辑。
"""

import asyncio
import tempfile
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from services.nlp import (
    ChapterSplitError,
    NovelText,
    RegexChapterRecognitionStrategy,
    validate_chapter_split,
)
from services.oss import oss


def read_plain_text(file_path: str) -> str:
    raw = Path(file_path).read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def extract_document_text(file_path: str, extension: str) -> str:
    """提取上传书稿正文，供 Agent 分章和模型分析。"""
    extension = extension.lower()
    if extension in {".txt", ".md"}:
        return read_plain_text(file_path)
    if extension == ".docx":
        document = Document(file_path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            paragraphs.extend(
                "\t".join(cell.text for cell in row.cells)
                for row in table.rows
            )
        return "\n".join(text for text in paragraphs if text.strip())
    if extension == ".pdf":
        reader = PdfReader(file_path)
        return "\n\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()
    return ""


def analyze_document(file_path: str, original_filename: str) -> dict:
    """提取正文并做章节结构校验（上传与 OSS 终局共用）。"""
    extension = Path(original_filename).suffix
    text_content = extract_document_text(file_path, extension)
    chapter_validation = None
    if text_content.strip():
        parsed_chapters = RegexChapterRecognitionStrategy().recognize(
            NovelText.from_string(text_content)
        )
        try:
            validate_chapter_split(text_content, parsed_chapters)
            chapter_validation = {
                "valid": True,
                "chapter_count": len(parsed_chapters) or 1,
                "text_length": len(text_content),
                "message": "书稿结构检查通过",
            }
        except ChapterSplitError as error:
            chapter_validation = {
                "valid": False,
                "chapter_count": len(parsed_chapters),
                "text_length": len(text_content),
                "message": str(error),
            }
    return {"text_content": text_content, "chapter_validation": chapter_validation}


async def analyze_oss_document(key: str, original_filename: str) -> dict:
    """经内网读取 OSS 对象并解析正文（不经过浏览器中转）。"""
    data = await oss.get_bytes(key)
    suffix = Path(original_filename).suffix
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as handle:
        handle.write(data)
        temp_path = handle.name
    try:
        return await asyncio.to_thread(analyze_document, temp_path, original_filename)
    finally:
        Path(temp_path).unlink(missing_ok=True)
