import pytest
import os
import tempfile
import shutil
from io import BytesIO

from docx import Document
from httpx import AsyncClient
from config import settings


@pytest.fixture(autouse=True)
def _ensure_media_dir(tmp_path):
    """使用临时目录替代真实 MEDIA_PATH，测试后自动清理。"""
    original = settings.MEDIA_PATH
    settings.MEDIA_PATH = str(tmp_path)
    yield
    settings.MEDIA_PATH = original


@pytest.mark.asyncio
async def test_upload_single_file(client: AsyncClient):
    """上传单个文件成功。"""
    file_content = b"hello world"
    response = await client.post(
        "/api/file/upload",
        files=[("files", ("test.txt", file_content, "text/plain"))],
    )
    assert response.status_code == 200, response.text
    data = response.json()["data"]
    assert data["total"] == 1
    assert data["files"][0]["original_filename"] == "test.txt"
    assert data["files"][0]["content_type"] == "text/plain"
    assert os.path.exists(data["files"][0]["file_path"])
    assert data["files"][0]["text_content"] == "hello world"
    print(f"    上传单文件: filename='{data['files'][0]['filename']}', 原始名='test.txt'")


@pytest.mark.asyncio
async def test_upload_docx_extracts_book_text(client: AsyncClient):
    """DOCX 上传后返回可供 Agent 分析的正文。"""
    document = Document()
    document.add_heading("第一章 山雨", level=1)
    document.add_paragraph("林舟在雨夜离开故乡。")
    content = BytesIO()
    document.save(content)

    response = await client.post(
        "/api/file/upload",
        files=[
            (
                "files",
                ("book.docx", content.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            )
        ],
    )

    assert response.status_code == 200, response.text
    text_content = response.json()["data"]["files"][0]["text_content"]
    assert "第一章 山雨" in text_content
    assert "林舟在雨夜离开故乡" in text_content


@pytest.mark.asyncio
async def test_upload_gb18030_text_previews_chapters(client: AsyncClient):
    """GB18030 小说应正确解码并在进入项目前识别章节数量。"""
    text = "第一章 山雨\n正文一\n第二章 归途\n正文二"
    response = await client.post(
        "/api/file/upload",
        files=[("files", ("gb-book.txt", text.encode("gb18030"), "text/plain"))],
    )

    assert response.status_code == 200, response.text
    uploaded = response.json()["data"]["files"][0]
    assert uploaded["text_content"] == text
    assert uploaded["chapter_validation"]["valid"] is True
    assert uploaded["chapter_validation"]["chapter_count"] == 2


@pytest.mark.asyncio
async def test_upload_long_unstructured_text_reports_invalid(client: AsyncClient):
    """长文本只识别到单章时，上传预检直接返回失败原因。"""
    text = "没有章节标题的长篇正文。" * 3_000
    response = await client.post(
        "/api/file/upload",
        files=[("files", ("invalid.txt", text.encode("utf-8"), "text/plain"))],
    )

    assert response.status_code == 200, response.text
    validation = response.json()["data"]["files"][0]["chapter_validation"]
    assert validation["valid"] is False
    assert validation["chapter_count"] == 0
    assert "只识别到 0 章" in validation["message"]


@pytest.mark.asyncio
async def test_upload_multiple_files(client: AsyncClient):
    """上传多个文件成功。"""
    response = await client.post(
        "/api/file/upload",
        files=[
            ("files", ("a.png", b"\x89PNG", "image/png")),
            ("files", ("b.jpg", b"\xff\xd8\xff", "image/jpeg")),
        ],
    )
    assert response.status_code == 200, response.text
    data = response.json()["data"]
    assert data["total"] == 2
    assert data["files"][0]["original_filename"] == "a.png"
    assert data["files"][1]["original_filename"] == "b.jpg"
    print(f"    上传多文件: total={data['total']}, files={[f['original_filename'] for f in data['files']]}")


@pytest.mark.asyncio
async def test_upload_long_filename_truncated(client: AsyncClient):
    """文件名超长时被截断到10个字符（保留扩展名）。"""
    long_name = "a" * 50 + ".txt"
    response = await client.post(
        "/api/file/upload",
        files=[("files", (long_name, b"data", "text/plain"))],
    )
    assert response.status_code == 200, response.text
    data = response.json()["data"]
    saved_filename = data["files"][0]["filename"]
    # 文件名中截断部分不应超过10个字符 (去掉时间戳前缀和扩展名)
    name_part = saved_filename.rsplit("_", 1)[-1]  # 取最后一个 _xxx.txt
    name_without_ext = os.path.splitext(name_part)[0]
    assert len(name_without_ext) <= 10
    print(f"    长文件名截断: 原始名='{long_name}' -> 保存为='{saved_filename}'")


@pytest.mark.asyncio
async def test_upload_to_invalid_path_returns_500(client: AsyncClient):
    """MEDIA_PATH 无效时返回 500 错误。"""
    settings.MEDIA_PATH = "/nonexistent/path/that/does/not/exist"
    response = await client.post(
        "/api/file/upload",
        files=[("files", ("fail.txt", b"data", "text/plain"))],
    )
    body = response.json()
    assert body["code"] == 500
    assert "上传失败" in body["message"]
    print(f"    无效路径上传失败: code={body['code']}, message='{body['message']}'")
