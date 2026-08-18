from fastapi import APIRouter, Depends, File, Query, UploadFile
from datetime import datetime
import asyncio
import os
import shutil
import tempfile
from pathlib import Path

from pydantic import BaseModel, Field

from docx import Document
from pypdf import PdfReader

from auth.deps import AuthContext, require_roles
from services.oss import make_upload_key, oss
from utils.response_format import ResponseSchema
from config import settings
from services.nlp import (
    ChapterSplitError,
    NovelText,
    RegexChapterRecognitionStrategy,
    validate_chapter_split,
)

router = APIRouter()

_EDITOR = Depends(require_roles("admin", "creator"))


class OssFinalizeIn(BaseModel):
    key: str = Field(min_length=1, max_length=500, description="OSS 对象 key")
    original_filename: str = Field(min_length=1, max_length=255, description="原始文件名")


@router.get("/upload-policy", summary="获取直传 OSS 的策略（未启用时 direct=false）")
async def get_upload_policy(
    filename: str = Query(..., max_length=255),
    content_type: str = Query("application/octet-stream", max_length=120),
    _: AuthContext = _EDITOR,
):
    if not oss.enabled:
        return ResponseSchema(data={"direct": False})
    key = make_upload_key(None, filename)
    policy = oss.sign_form_upload(key, content_type, 20 * 1024 * 1024)
    return ResponseSchema(
        data={
            "direct": True,
            "provider": oss.name,
            "key": key,
            "upload_url": policy["url"],
            "fields": policy["fields"],
            "public_url": oss.public_url(key),
            "filename": Path(filename).name,
        }
    )


@router.post("/oss-finalize", summary="OSS 直传后终局：服务端经内网读取并做书稿分析")
async def oss_finalize(
    payload: OssFinalizeIn,
    _: AuthContext = _EDITOR,
):
    if not oss.enabled:
        from fastapi import HTTPException

        raise HTTPException(status_code=400, detail="未启用对象存储")
    data = await oss.get_bytes(payload.key)
    with tempfile.NamedTemporaryFile(suffix=Path(payload.original_filename).suffix, delete=False) as handle:
        handle.write(data)
        temp_path = handle.name
    try:
        analysis = await asyncio.to_thread(
            _analyze_document, temp_path, payload.original_filename
        )
    finally:
        Path(temp_path).unlink(missing_ok=True)
    return ResponseSchema(
        data={
            "filename": Path(payload.original_filename).name,
            "url": oss.public_url(payload.key),
            "key": payload.key,
            "text_content": analysis["text_content"],
            "chapter_validation": analysis["chapter_validation"],
            "message": "文件上传成功",
        }
    )


def _analyze_document(file_path: str, original_filename: str) -> dict:
    """提取正文并做章节结构校验（上传与 OSS 终局共用）。"""
    extension = Path(original_filename).suffix
    text_content = _extract_document_text(file_path, extension)
    chapter_validation = None
    if text_content.strip():
        parsed_chapters = RegexChapterRecognitionStrategy().recognize(
            NovelText.from_string(text_content)
        )
        try:
            validate_chapter_split(text_content, parsed_chapters)
            chapter_validation = {
                "valid": True,
                "chapter_count": len(parsed_chapters) or 1,
                "text_length": len(text_content),
                "message": "书稿结构检查通过",
            }
        except ChapterSplitError as error:
            chapter_validation = {
                "valid": False,
                "chapter_count": len(parsed_chapters),
                "text_length": len(text_content),
                "message": str(error),
            }
    return {"text_content": text_content, "chapter_validation": chapter_validation}


def _read_plain_text(file_path: str) -> str:
    raw = Path(file_path).read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _extract_document_text(file_path: str, extension: str) -> str:
    """提取上传书稿正文，供 Agent 分章和模型分析。"""
    extension = extension.lower()
    if extension in {".txt", ".md"}:
        return _read_plain_text(file_path)
    if extension == ".docx":
        document = Document(file_path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            paragraphs.extend(
                "\t".join(cell.text for cell in row.cells)
                for row in table.rows
            )
        return "\n".join(text for text in paragraphs if text.strip())
    if extension == ".pdf":
        reader = PdfReader(file_path)
        return "\n\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()
    return ""

@router.post("/upload", summary="多文件上传", response_model=ResponseSchema[dict])
async def upload_files(
    files: list[UploadFile] = File(...),
    _: AuthContext = _EDITOR,
):
    """处理多文件上传"""
    try:
        results = []

        # 循环处理每个文件
        for file in files:
            # 生成唯一的文件名，避免覆盖
            timestamp = datetime.now(settings.tz).strftime("%Y%m%d%H%M%S")
            # 为每个文件添加唯一标识，避免同一时间上传的文件重名
            unique_id = f"{timestamp}_{os.urandom(4).hex()}"

            # 截断原始文件名到最长10个字符，但保留扩展名
            original_filename = file.filename
            # 分离文件名和扩展名
            name_part, ext_part = os.path.splitext(original_filename)
            # 只截取文件名部分的前10个字符，保留扩展名
            truncated_name = name_part[:10]
            # 构建新文件名（保留原始扩展名）
            filename = f"{unique_id}_{truncated_name}{ext_part}"
            file_path = os.path.join(settings.MEDIA_PATH, filename)

            # 保存文件
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)

            analysis = await asyncio.to_thread(
                _analyze_document,
                file_path,
                original_filename,
            )
            text_content = analysis["text_content"]
            chapter_validation = analysis["chapter_validation"]

            # 记录结果
            results.append(
                {
                    "filename": filename,
                    "original_filename": original_filename,
                    "content_type": file.content_type,
                    "file_path": file_path,
                    "text_content": text_content,
                    "chapter_validation": chapter_validation,
                    "message": "文件上传成功",
                }
            )

        return ResponseSchema(data={"total": len(results), "files": results})
    except Exception as e:
        return ResponseSchema(code=500, data={"error": str(e)}, message="上传失败")
